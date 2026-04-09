import os
from pathlib import Path
from typing import List, Dict
import json

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("python-docx not installed. Install with: pip install python-docx")

def convert_docx_to_text(docx_path: Path) -> str:
    """
    Convert a .docx file to structured text format.
    
    Args:
        docx_path: Path to the .docx file
    
    Returns:
        Structured text content
    """
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx package is required. Install with: pip install python-docx")
    
    try:
        doc = Document(docx_path)
        
        # Extract filename without extension for title
        filename = docx_path.stem
        
        # Start with structured format
        structured_content = f"# Meeting: {filename}\n\n"
        
        # Extract all paragraphs
        content_paragraphs = []
        current_section = None
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            
            if not text:
                continue
            
            # Check if paragraph looks like a heading (all caps, short, or starts with common meeting terms)
            if (is_likely_heading(text)):
                # Add as section header
                structured_content += f"## {text}\n\n"
                current_section = text
            else:
                # Add as regular content
                structured_content += f"{text}\n\n"
        
        return structured_content
        
    except Exception as e:
        print(f"Error processing {docx_path.name}: {e}")
        return ""

def is_likely_heading(text: str) -> bool:
    """
    Determine if text is likely a section heading.
    """
    # Common meeting section indicators
    heading_keywords = [
        'date', 'client', 'company', 'participants', 'attendees', 
        'summary', 'agenda', 'discussion', 'key points', 'action items',
        'next steps', 'follow up', 'objections', 'concerns', 'pricing',
        'timeline', 'budget', 'requirements', 'features', 'technical',
        'implementation', 'security', 'integration', 'competitors'
    ]
    
    text_lower = text.lower()
    
    # Check if it's short and contains heading keywords
    if len(text) < 50 and any(keyword in text_lower for keyword in heading_keywords):
        return True
    
    # Check if it's all caps (likely a heading)
    if text.isupper() and len(text) < 100:
        return True
    
    # Check if it ends with colon (common in meeting notes)
    if text.endswith(':') and len(text) < 100:
        return True
    
    return False

def convert_all_docx_files(
    input_dir: str = "data/raw",
    output_dir: str = "data/raw"
) -> List[str]:
    """
    Convert all .docx files in the input directory to .txt format.
    
    Args:
        input_dir: Directory containing .docx files
        output_dir: Directory to save converted .txt files
    
    Returns:
        List of converted filenames
    """
    if not DOCX_AVAILABLE:
        print("❌ python-docx not installed!")
        print("Install it with: pip install python-docx")
        return []
    
    input_path = Path(input_dir)
    output_path = Path(output_dir)
    
    if not input_path.exists():
        print(f"Input directory {input_dir} doesn't exist")
        return []
    
    output_path.mkdir(parents=True, exist_ok=True)
    
    docx_files = list(input_path.glob("*.docx"))
    
    if not docx_files:
        print(f"No .docx files found in {input_dir}")
        return []
    
    converted_files = []
    
    print(f"Found {len(docx_files)} .docx files to convert...")
    
    for docx_file in docx_files:
        print(f"Converting: {docx_file.name}")
        
        try:
            # Convert to text
            text_content = convert_docx_to_text(docx_file)
            
            if text_content:
                # Save as .txt file
                output_file = output_path / f"{docx_file.stem}.txt"
                
                with open(output_file, 'w', encoding='utf-8') as f:
                    f.write(text_content)
                
                converted_files.append(output_file.name)
                print(f"✅ Converted to: {output_file.name}")
            else:
                print(f"❌ Failed to extract content from: {docx_file.name}")
                
        except Exception as e:
            print(f"❌ Error converting {docx_file.name}: {e}")
    
    print(f"\n🎉 Successfully converted {len(converted_files)} files:")
    for filename in converted_files:
        print(f"   • {filename}")
    
    return converted_files

def extract_tables_from_docx(docx_path: Path) -> List[Dict]:
    """
    Extract tables from .docx file (useful for structured meeting data).
    
    Args:
        docx_path: Path to the .docx file
    
    Returns:
        List of table data
    """
    if not DOCX_AVAILABLE:
        return []
    
    try:
        doc = Document(docx_path)
        tables_data = []
        
        for i, table in enumerate(doc.tables):
            table_data = {
                'table_index': i,
                'rows': []
            }
            
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                if any(row_data):  # Only add non-empty rows
                    table_data['rows'].append(row_data)
            
            if table_data['rows']:
                tables_data.append(table_data)
        
        return tables_data
        
    except Exception as e:
        print(f"Error extracting tables from {docx_path.name}: {e}")
        return []

def create_sample_meeting_files():
    """
    Create sample meeting files in different formats for testing.
    """
    sample_dir = Path("data/raw/samples")
    sample_dir.mkdir(parents=True, exist_ok=True)
    
    # Sample JSON format
    sample_json = {
        "title": "Enterprise Client Discovery Call - TechCorp",
        "date": "2024-01-15",
        "client": "TechCorp Inc",
        "participants": ["John Smith (Sales)", "Jane Doe (CTO)", "Mike Johnson (Technical Lead)"],
        "meeting_type": "Discovery Call",
        "summary": "Initial discovery call with enterprise client to understand their delivery platform needs. They're looking to launch a multi-vendor marketplace similar to Uber Eats but for their specific industry vertical.",
        "key_points": [
            "Budget range: $75k-150k for complete solution",
            "Timeline: Q2 2024 launch target",
            "Main concern: scalability and performance",
            "Need white-label solution with custom branding",
            "Require integration with existing payment systems"
        ],
        "objections": "Concerned about development timeline and whether our solution can handle their expected volume of 10k+ orders per day",
        "pricing_discussion": "Discussed Enterprise tier at $120k. Client interested but wants to see performance benchmarks and case studies from similar scale deployments.",
        "competitor_mentions": "Currently evaluating our solution against custom development and another vendor (FoodTech Solutions)",
        "technical_requirements": [
            "Multi-tenant architecture",
            "API integrations with existing systems",
            "Custom payment gateway integration",
            "Advanced analytics and reporting"
        ],
        "next_steps": [
            "Send performance benchmarks and case studies",
            "Schedule technical demo with their development team",
            "Provide detailed implementation timeline",
            "Connect them with existing enterprise client for reference"
        ],
        "follow_up": "Follow up in 3 days with requested materials. Schedule technical demo for next week."
    }
    
    with open(sample_dir / "sample_meeting.json", 'w', encoding='utf-8') as f:
        json.dump(sample_json, f, indent=2, ensure_ascii=False)
    
    # Sample text format
    sample_text = """# Meeting: SMB Client Consultation - LocalEats

## Date
2024-01-20

## Client
LocalEats Restaurant Group

## Participants
Sarah Wilson (Sales), Tom Brown (Owner), Lisa Chen (Operations Manager)

## Meeting Type
Consultation Call

## Summary
Consultation with local restaurant group looking to add delivery capabilities to their 5 restaurant locations. They want to start with delivery and potentially expand to third-party restaurants later.

## Key Points
• Budget: $25k-40k range
• Timeline: 6-8 weeks implementation
• Start with single-vendor, expand to multi-vendor later
• Focus on local delivery radius (5-mile radius)
• Need integration with existing POS system

## Client Concerns
• Worried about complexity of managing delivery operations
• Concerned about driver management and reliability
• Questions about commission structure for future third-party restaurants

## Pricing Discussion
Recommended Single Vendor solution at $35k. Client very interested and within budget. Discussed optional driver management service for additional $500/month.

## Technical Requirements
• POS integration (they use Square)
• Simple admin dashboard
• Customer mobile app with their branding
• Basic analytics and reporting

## Competitor Analysis
They mentioned looking at DoorDash for Business but prefer owning their own platform for better control and customer data.

## Next Steps
• Send proposal with Single Vendor package details
• Schedule demo of customer app and admin dashboard
• Provide case study from similar restaurant group
• Discuss implementation timeline and milestones

## Follow-up
Send proposal by end of week. They want to make decision by month-end.
"""
    
    with open(sample_dir / "sample_meeting.txt", 'w', encoding='utf-8') as f:
        f.write(sample_text)
    
    print(f"✅ Created sample meeting files in {sample_dir}")
    print("   • sample_meeting.json")
    print("   • sample_meeting.txt")

if __name__ == "__main__":
    import sys
    
    if len(sys.argv) > 1:
        command = sys.argv[1].lower()
        
        if command == "convert":
            # Convert all .docx files
            converted = convert_all_docx_files()
            if converted:
                print(f"\n🎯 Next steps:")
                print("1. Review the converted .txt files")
                print("2. Run the pipeline: ./run_pipeline.sh")
                print("3. Start the chatbot: uvicorn api.main:app --reload --port 8000")
        
        elif command == "samples":
            # Create sample files
            create_sample_meeting_files()
        
        else:
            print("Usage:")
            print("  python convert_docx.py convert  - Convert all .docx files to .txt")
            print("  python convert_docx.py samples  - Create sample meeting files")
    
    else:
        print("DOCX Converter for Sales Meeting Chatbot")
        print("=" * 40)
        print()
        print("Usage:")
        print("  python convert_docx.py convert  - Convert all .docx files in data/raw/")
        print("  python convert_docx.py samples  - Create sample meeting files")
        print()
        print("Requirements:")
        print("  pip install python-docx")